"""
contract_analyzer_vsc.py

Single-shot Contract Simplifier + Bias Detector (for VS Code / local use).

- Uses OpenAI API (OPENAI_API_KEY env var) to perform:
    * whole-document simplification (plain + bullets)
    * overall bias detection (Low/Moderate/High) with reasons
    * a fair rewrite (excerpt)
- Saves outputs to TXT, CSV, DOCX and prints a short console summary.
- If OpenAI key is not present, falls back to simple heuristics (keyword-based bias + no-op simplifier).

Requires:
    pip install openai pdfplumber nltk python-docx pandas
    python -c "import nltk; nltk.download('punkt')"
"""
from __future__ import annotations
import os
import sys
import json
import re
import argparse
import tempfile
import io
import time
import csv
from typing import Tuple, Dict, Any

# optional libs
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import openai
except Exception:
    openai = None

try:
    import pandas as pd
except Exception:
    pd = None

try:
    from docx import Document
except Exception:
    Document = None

import nltk
nltk.download('punkt', quiet=True)

# ---------- CONFIG ----------
DEFAULT_INPUT_PDF = "contract.pdf"
DEFAULT_INPUT_TXT = "contract.txt"
OPENAI_ENV_VAR = "OPENAI_API_KEY"
DEFAULT_OPENAI_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")  # adjustable

BIAS_KEYWORDS = [
    "sole discretion", "at its discretion", "without notice", "indemnify", "irrevocable",
    "perpetual", "forever", "not liable", "without liability", "no refund", "non-refundable",
    "no recourse", "exclusive", "only the", "anytime", "at will", "terminate for any reason",
    "binding on", "assigns and successors", "waive", "waives", "waiver", "no appeal",
    "court waived", "arbitration", "no jury", "liquidated damages", "liability"
]

# ---------- utilities ----------
def simple_clean(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\r", " ").strip()
    # normalize whitespace but keep paragraph breaks
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()

def extract_text_from_pdf_path(path: str) -> str:
    if pdfplumber is None:
        raise RuntimeError("pdfplumber not installed. Install with: pip install pdfplumber")
    pages = []
    with pdfplumber.open(path) as pdf:
        for p in pdf.pages:
            pages.append(p.extract_text() or "")
    return "\n\n".join(pages)

def read_input_text(path: str | None) -> Tuple[str, str]:
    """
    Returns (raw_text, source_label)
    """
    if path:
        p = path
        if not os.path.exists(p):
            raise FileNotFoundError(p)
        if p.lower().endswith(".pdf"):
            raw = extract_text_from_pdf_path(p)
            return raw, p
        else:
            raw = open(p, "r", encoding="utf-8", errors="ignore").read()
            return raw, p
    # auto mode: prefer PDF then TXT
    if os.path.exists(DEFAULT_INPUT_PDF):
        return extract_text_from_pdf_path(DEFAULT_INPUT_PDF), DEFAULT_INPUT_PDF
    if os.path.exists(DEFAULT_INPUT_TXT):
        return open(DEFAULT_INPUT_TXT, "r", encoding="utf-8", errors="ignore").read(), DEFAULT_INPUT_TXT
    raise FileNotFoundError("No input file found. Provide --file PATH or place contract.pdf / contract.txt in folder.")

# ---------- OpenAI helpers (single-shot prompt) ----------
def call_openai_json(prompt_instruction: str, model: str = DEFAULT_OPENAI_MODEL, api_key: str | None = None, timeout: int = 300) -> Dict[str, Any]:
    """
    Calls OpenAI ChatCompletion (Chat API) with system->user instructing JSON-only response.
    Expects the model to return a single JSON object. Defensive parsing included.
    """
    if openai is None:
        raise RuntimeError("openai package not installed. pip install openai")

    key = api_key or os.environ.get(OPENAI_ENV_VAR)
    if not key:
        raise RuntimeError(f"OpenAI API key not found. Set env var {OPENAI_ENV_VAR} or pass api_key.")
    openai.api_key = key

    system = (
        "You are a contract assistant. Output ONLY valid JSON, no surrounding text. "
        "Fields required in the JSON: simplified_plain, simplified_bullets, bias_label, bias_score, bias_reasons, fair_rewrite_excerpt. "
        "Use concise values. bias_label must be one of: Low, Moderate, High."
    )

    messages = [
        {"role":"system", "content": system},
        {"role":"user", "content": prompt_instruction}
    ]

    # Use ChatCompletion if available
    try:
        resp = openai.ChatCompletion.create(
            model=model,
            messages=messages,
            temperature=0.0,
            max_tokens=1500,
            timeout=timeout
        )
        text = resp["choices"][0]["message"]["content"]
    except Exception as e:
        # try a simpler completion (text-davinci) fallback
        try:
            resp = openai.Completion.create(engine="text-davinci-003", prompt=system + "\n\n" + prompt_instruction, max_tokens=1500, temperature=0.0)
            text = resp["choices"][0]["text"]
        except Exception as ee:
            raise RuntimeError(f"OpenAI call failed: {e} / fallback: {ee}")

    # defensive: extract JSON substring from response
    json_text = extract_json_text(text)
    if not json_text:
        raise ValueError("OpenAI response did not contain valid JSON. Raw response:\n" + text)
    return json.loads(json_text)

def extract_json_text(s: str) -> str | None:
    """
    Finds the first JSON object substring in a text.
    """
    s = s.strip()
    # naive scan for matching braces
    start = s.find("{")
    if start == -1:
        return None
    depth = 0
    for i in range(start, len(s)):
        if s[i] == "{":
            depth += 1
        elif s[i] == "}":
            depth -= 1
            if depth == 0:
                return s[start:i+1]
    return None

# ---------- Local fallback heuristics ----------
def local_detect_bias(text: str) -> Dict[str, Any]:
    """
    Simple keyword-based bias detection (fallback).
    """
    reasons = []
    matched = []
    t = text.lower()
    hits = 0
    for kw in BIAS_KEYWORDS:
        if kw in t:
            hits += 1
            matched.append(kw)
            reasons.append(f"Contains phrase: '{kw}'")
    score = min(0.15 * hits, 0.9)
    if "sole discretion" in t or "without notice" in t:
        score = min(1.0, score + 0.2)
    if score < 0.25:
        label = "Low"
    elif score < 0.55:
        label = "Moderate"
    else:
        label = "High"
    return {"bias_label": label, "bias_score": round(score, 3), "bias_reasons": reasons, "matched_keywords": matched}

def local_simplify(text: str) -> Tuple[str, str]:
    """
    Very naive fallback: split into short sentences and join into bullets.
    """
    import nltk
    sents = nltk.tokenize.sent_tokenize(text)
    plain = " ".join(sents[:10]) + (" ..." if len(sents) > 10 else "")
    bullets = "\n- " + "\n- ".join([s.strip() for s in sents[:8]])
    return plain, bullets

def local_fair_rewrite(text: str) -> str:
    s = text
    s = re.sub(r"at (its|their) sole discretion", "after mutual agreement or reasonable notice", s, flags=re.IGNORECASE)
    s = re.sub(r"without notice", "with reasonable prior notice", s, flags=re.IGNORECASE)
    return s

# ---------- Main pipeline ----------
def analyze_contract_text(raw_text: str, use_openai: bool = True, openai_model: str = DEFAULT_OPENAI_MODEL, api_key: str | None = None, simplify_style: str = "both") -> Dict[str, Any]:
    t = simple_clean(raw_text)
    if not t:
        raise ValueError("Empty contract text")

    # If OpenAI available and key present, call single-shot JSON prompt
    if use_openai and openai is not None and (api_key or os.environ.get(OPENAI_ENV_VAR)):
        # guard length: for enormous documents, pass head+tail
        if len(t) > 32000:
            t_pass = t[:16000] + "\n\n" + t[-16000:]
        else:
            t_pass = t

        prompt = (
            "Simplify the following contract for a non-lawyer and flag overall bias. "
            "Return JSON with keys: simplified_plain (string), simplified_bullets (string), "
            "bias_label (Low|Moderate|High), bias_score (0..1 float), bias_reasons (list of strings), "
            "fair_rewrite_excerpt (string). Keep values concise.\n\n"
            "Contract Text:\n" + t_pass + "\n\n"
            "Produce JSON only."
        )
        try:
            resp = call_openai_json(prompt, model=openai_model, api_key=api_key)
            # Ensure fields present; if not, fallback to local
            required = {"simplified_plain","simplified_bullets","bias_label","bias_score","bias_reasons","fair_rewrite_excerpt"}
            if not required.issubset(set(resp.keys())):
                # try to normalize if model used different keys
                # minimal normalization attempts
                resp2 = {}
                resp2["simplified_plain"] = resp.get("plain", resp.get("simplified_plain", ""))
                resp2["simplified_bullets"] = resp.get("bullets", resp.get("simplified_bullets", ""))
                resp2["bias_label"] = resp.get("bias_label", resp.get("label", "Unknown"))
                resp2["bias_score"] = resp.get("bias_score", resp.get("score", 0.0))
                resp2["bias_reasons"] = resp.get("bias_reasons", resp.get("reasons", []))
                resp2["fair_rewrite_excerpt"] = resp.get("fair_rewrite_excerpt", resp.get("fair_rewrite", ""))
                resp = resp2
            # finally attach original excerpt
            resp["original_excerpt"] = t[:20000] + ("..." if len(t) > 20000 else "")
            return resp
        except Exception as e:
            # Fall back to local processing and record the error
            print("OpenAI call failed; falling back to local heuristics. Error:", str(e))
            local_plain, local_bullets = local_simplify(t)
            lbias = local_detect_bias(t)
            fair = local_fair_rewrite(t)
            return {
                "simplified_plain": local_plain,
                "simplified_bullets": local_bullets,
                "bias_label": lbias["bias_label"],
                "bias_score": lbias["bias_score"],
                "bias_reasons": lbias["bias_reasons"],
                "fair_rewrite_excerpt": fair,
                "original_excerpt": t[:20000] + ("..." if len(t) > 20000 else "")
            }

    # if OpenAI is not configured, do local fallback
    local_plain, local_bullets = local_simplify(t)
    lbias = local_detect_bias(t)
    fair = local_fair_rewrite(t)
    return {
        "simplified_plain": local_plain,
        "simplified_bullets": local_bullets,
        "bias_label": lbias["bias_label"],
        "bias_score": lbias["bias_score"],
        "bias_reasons": lbias["bias_reasons"],
        "fair_rewrite_excerpt": fair,
        "original_excerpt": t[:20000] + ("..." if len(t) > 20000 else "")
    }

# ---------- Save outputs ----------
def save_outputs(resp: Dict[str, Any], out_prefix: str = "contract_analysis"):
    txt_path = out_prefix + ".txt"
    csv_path = out_prefix + ".csv"
    docx_path = out_prefix + ".docx"

    # TXT
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("=== ORIGINAL EXCERPT ===\n\n")
        f.write(resp.get("original_excerpt", "") + "\n\n")
        f.write("=== SIMPLIFIED (PLAIN) ===\n\n")
        f.write(resp.get("simplified_plain", "") + "\n\n")
        f.write("=== SIMPLIFIED (BULLETS) ===\n\n")
        f.write(resp.get("simplified_bullets", "") + "\n\n")
        f.write("=== BIAS ===\n\n")
        f.write(f"Label: {resp.get('bias_label')} (score: {resp.get('bias_score')})\n")
        f.write("Reasons:\n")
        for r in resp.get("bias_reasons", []) or []:
            f.write(" - " + str(r) + "\n")
        f.write("\n=== FAIR REWRITE (EXCERPT) ===\n\n")
        f.write(resp.get("fair_rewrite_excerpt", "") + "\n")

    # CSV (single row)
    row = {
        "original_excerpt": resp.get("original_excerpt", ""),
        "simplified_plain": resp.get("simplified_plain", ""),
        "simplified_bullets": resp.get("simplified_bullets", ""),
        "bias_label": resp.get("bias_label", ""),
        "bias_score": resp.get("bias_score", ""),
        "bias_reasons": "; ".join(resp.get("bias_reasons", []) or []),
        "fair_rewrite_excerpt": resp.get("fair_rewrite_excerpt", "")
    }
    if pd is not None:
        df = pd.DataFrame([row])
        df.to_csv(csv_path, index=False, encoding="utf-8")
    else:
        with open(csv_path, "w", encoding="utf-8", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=list(row.keys()))
            writer.writeheader()
            writer.writerow(row)

    # DOCX
    if Document is not None:
        doc = Document()
        doc.add_heading("Contract Analysis", level=1)
        doc.add_heading("Original Excerpt", level=2)
        doc.add_paragraph(resp.get("original_excerpt", ""))
        doc.add_heading("Simplified (Plain)", level=2)
        doc.add_paragraph(resp.get("simplified_plain", ""))
        doc.add_heading("Simplified (Bullets)", level=2)
        doc.add_paragraph(resp.get("simplified_bullets", ""))
        doc.add_heading("Bias", level=2)
        doc.add_paragraph(f"Label: {resp.get('bias_label')} (score: {resp.get('bias_score')})")
        for r in resp.get("bias_reasons", []) or []:
            doc.add_paragraph("- " + str(r))
        doc.add_heading("Fair Rewrite (Excerpt)", level=2)
        doc.add_paragraph(resp.get("fair_rewrite_excerpt", ""))
        doc.save(docx_path)
    else:
        # if python-docx not installed, create a simple .txt copy named .docx.txt (informational)
        with open(docx_path + ".txt", "w", encoding="utf-8") as f:
            f.write("python-docx not installed. Install 'python-docx' to generate a real .docx file.\n\n")
            f.write("Original Excerpt:\n")
            f.write(resp.get("original_excerpt", "") + "\n\n")
            f.write("Simplified (Plain):\n")
            f.write(resp.get("simplified_plain", "") + "\n\n")

    return {"txt": txt_path, "csv": csv_path, "docx": docx_path}

# ---------- CLI ----------
def main_cli():
    ap = argparse.ArgumentParser(description="Contract Simplifier + Bias Detector (single-shot).")
    ap.add_argument("--file", "-f", type=str, help="Path to contract PDF or TXT. If omitted, script will look for contract.pdf or contract.txt in current folder.")
    ap.add_argument("--no-openai", action="store_true", help="Do not call OpenAI; use local heuristics only.")
    ap.add_argument("--model", type=str, default=DEFAULT_OPENAI_MODEL, help="OpenAI model to call (optional).")
    ap.add_argument("--out-prefix", type=str, default="contract_analysis", help="Output file prefix (contract_analysis by default).")
    args = ap.parse_args()

    try:
        raw_text, source = read_input_text(args.file)
    except Exception as e:
        print("Error reading input:", e)
        sys.exit(1)

    use_openai = (not args.no_openai) and (openai is not None) and (os.environ.get(OPENAI_ENV_VAR) is not None)
    if not use_openai:
        print("OpenAI not used (either --no-openai, openai package missing, or OPENAI_API_KEY not set). Using local heuristics.")

    print("Analyzing (this may take a few seconds)...")
    start = time.time()
    resp = analyze_contract_text(raw_text, use_openai=use_openai, openai_model=args.model, api_key=os.environ.get(OPENAI_ENV_VAR))
    took = time.time() - start

    print("\n=== SUMMARY ===")
    print(f"Source: {source}")
    print(f"Elapsed: {took:.2f}s")
    print(f"Bias: {resp.get('bias_label')} (score: {resp.get('bias_score')})")
    print("\nSimplified (plain) excerpt:\n")
    sp = resp.get("simplified_plain", "")
    print(sp[:1000] + ("..." if len(sp) > 1000 else ""))
    print("\nSaving outputs...")
    outpaths = save_outputs(resp, out_prefix=args.out_prefix)
    print("Saved:", outpaths)
    print("Done.")

if __name__ == "__main__":
    main_cli()
